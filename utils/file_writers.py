
"""
File Writers - Generate normalized output files preserving original format.
Theme aligned with the Streamlit UI.
"""

import io
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_CENTER
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn

PRIMARY = "#0B7A75"
PRIMARY_DARK = "#084C61"
BORDER = "#CFE7E5"
ROW_ALT = "#F4FBFB"
TEXT_DARK = "#14323A"


def _records_to_output_df(records: list[dict]) -> pd.DataFrame:
    df = pd.DataFrame(records)

    if df.empty:
        return pd.DataFrame(columns=["Output Attribute", "Value"])

    rename_map = {}
    if "attribute" in df.columns:
        rename_map["attribute"] = "Output Attribute"
    if "value" in df.columns:
        rename_map["value"] = "Value"
    df = df.rename(columns=rename_map)

    preferred = []
    for col in ["Output Attribute", "Value"]:
        if col in df.columns:
            preferred.append(col)

    remaining = [c for c in df.columns if c not in preferred]
    return df[preferred + remaining]


# ── PDF WRITER ────────────────────────────────────────────────────────────────

def write_pdf_keyvalue(records: list[dict], title: str = "Normalized Document") -> bytes:
    """
    For PDF, keep a readable 2-column business output.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle(
        "title",
        parent=styles["Title"],
        fontSize=16,
        spaceAfter=12,
        alignment=TA_CENTER,
        textColor=colors.HexColor(PRIMARY_DARK),
    )
    subtitle_style = ParagraphStyle(
        "subtitle",
        parent=styles["Normal"],
        fontSize=9,
        alignment=TA_CENTER,
        textColor=colors.HexColor(TEXT_DARK),
        spaceAfter=10,
    )
    story.append(Paragraph(title, title_style))
    story.append(Paragraph("Normalized output with canonical attribute names and original values", subtitle_style))
    story.append(Spacer(1, 0.3 * cm))

    header_style = ParagraphStyle(
        "hdr",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Helvetica-Bold",
        textColor=colors.white,
    )
    cell_style = ParagraphStyle(
        "cell",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        textColor=colors.HexColor(TEXT_DARK),
    )

    table_data = [[Paragraph("Output Attribute", header_style), Paragraph("Value", header_style)]]
    for rec in records:
        table_data.append(
            [
                Paragraph(str(rec.get("attribute", "")), cell_style),
                Paragraph(str(rec.get("value", "")), cell_style),
            ]
        )

    table = Table(table_data, colWidths=[7.5 * cm, 10.5 * cm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(PRIMARY_DARK)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor(ROW_ALT), colors.white]),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor(BORDER)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.append(table)
    doc.build(story)
    return buf.getvalue()


def write_pdf_tabular(dfs: list[pd.DataFrame], title: str = "Normalized Document") -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=1 * cm,
        leftMargin=1 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle(
        "title",
        parent=styles["Title"],
        fontSize=14,
        spaceAfter=10,
        alignment=TA_CENTER,
        textColor=colors.HexColor(PRIMARY_DARK),
    )
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.3 * cm))

    for i, df in enumerate(dfs):
        if df.empty:
            continue

        if len(dfs) > 1:
            story.append(Paragraph(f"Table {i + 1}", styles["Heading2"]))

        col_style = ParagraphStyle(
            "col",
            parent=styles["Normal"],
            fontSize=8,
            fontName="Helvetica-Bold",
            textColor=colors.white,
        )
        cell_style = ParagraphStyle(
            "cell",
            parent=styles["Normal"],
            fontSize=8,
            leading=11,
            textColor=colors.HexColor(TEXT_DARK),
        )

        table_data = [[Paragraph(str(c), col_style) for c in df.columns]]
        for _, row in df.iterrows():
            table_data.append([Paragraph(str(v), cell_style) for v in row])

        n_cols = max(len(df.columns), 1)
        avail = 19 * cm
        table = Table(table_data, colWidths=[avail / n_cols] * n_cols, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(PRIMARY_DARK)),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor(ROW_ALT), colors.white]),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor(BORDER)),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 0.4 * cm))

    doc.build(story)
    return buf.getvalue()


def write_docx_keyvalue(records: list[dict], title: str = "Normalized Document") -> bytes:
    """
    Write normalized output to a DOCX file.
    Keeps Output Attribute and Value as the first two columns.
    If extra columns exist in records, they are also preserved.
    """
    df = _records_to_output_df(records)

    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(8, 76, 97)   # PRIMARY_DARK

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run("Normalized output with canonical attribute names and original values")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(20, 50, 58)

    doc.add_paragraph("")

    # Table
    rows = len(df) + 1
    cols = len(df.columns)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for col_idx, col_name in enumerate(df.columns):
        cell = table.cell(0, col_idx)
        cell.text = str(col_name)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)

        # Header shading
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.xpath('./w:shd')
        if not shd:
            from docx.oxml import OxmlElement
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '084C61')   # PRIMARY_DARK
            tc_pr.append(shd)

    # Data rows
    for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = "" if pd.isna(value) else str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if str(df.columns[col_idx]) == "Output Attribute":
                        run.bold = True
                        run.font.color.rgb = RGBColor(8, 76, 97)

    # Optional: set rough column widths
    try:
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx == 0:
                    cell.width = Inches(2.8)
                elif idx == 1:
                    cell.width = Inches(4.5)
                else:
                    cell.width = Inches(2.5)
    except Exception:
        pass

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ── EXCEL WRITER ──────────────────────────────────────────────────────────────

def _header_style():
    return {
        "font": Font(bold=True, color="FFFFFF", size=11),
        "fill": PatternFill("solid", start_color=PRIMARY_DARK.replace("#", "")),
        "alignment": Alignment(horizontal="center", vertical="center", wrap_text=True),
    }


def _apply(cell, **kwargs):
    for k, v in kwargs.items():
        setattr(cell, k, v)


def write_excel_keyvalue(records: list[dict], sheet_name: str = "Normalized") -> bytes:
    """
    Now supports dynamic multi-column records.
    Keeps Output Attribute and Value first, then preserves remaining columns.
    """
    df = _records_to_output_df(records)
    return write_excel_tabular({sheet_name: df})


def write_excel_tabular(sheets: dict[str, pd.DataFrame]) -> bytes:
    wb = Workbook()
    wb.remove(wb.active)

    for sheet_name, df in sheets.items():
        ws = wb.create_sheet(title=sheet_name[:31])
        hdr = _header_style()
        thin = Border(
            left=Side(style="thin", color=BORDER.replace("#", "")),
            right=Side(style="thin", color=BORDER.replace("#", "")),
            top=Side(style="thin", color=BORDER.replace("#", "")),
            bottom=Side(style="thin", color=BORDER.replace("#", "")),
        )

        for col_idx, col_name in enumerate(df.columns, start=1):
            c = ws.cell(row=1, column=col_idx, value=col_name)
            _apply(c, **hdr)

        for r_idx, (_, row) in enumerate(df.iterrows(), start=2):
            fill_color = ROW_ALT.replace("#", "") if r_idx % 2 == 0 else "FFFFFF"
            fill = PatternFill("solid", start_color=fill_color)
            for col_idx, val in enumerate(row, start=1):
                c = ws.cell(row=r_idx, column=col_idx, value=val)
                c.border = thin
                c.fill = fill
                c.alignment = Alignment(vertical="top", wrap_text=True)
                if str(df.columns[col_idx - 1]) == "Output Attribute":
                    c.font = Font(bold=True, color=PRIMARY_DARK.replace("#", ""))

        for col_idx, col in enumerate(df.columns, start=1):
            lengths = [len(str(col))]
            if len(df):
                lengths.extend(len(str(v)) for v in df.iloc[:, col_idx - 1].tolist())
            max_len = max(lengths) if lengths else 10
            ws.column_dimensions[ws.cell(1, col_idx).column_letter].width = min(max_len + 4, 45)

        ws.freeze_panes = "A2"
        ws.row_dimensions[1].height = 24

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── CSV WRITER ────────────────────────────────────────────────────────────────

def write_csv(df: pd.DataFrame) -> bytes:
    buf = io.StringIO()
    df.to_csv(buf, index=False)
    return buf.getvalue().encode("utf-8")


def write_csv_records(records: list[dict]) -> bytes:
    df = _records_to_output_df(records)
    return write_csv(df)