
"""
File Parsers - Extract attribute-value pairs from PDF, Excel, and CSV files.

Behavior for this project:
- Prefer clean key-value extraction for normalization output.
- Even when source looks tabular, if rows behave like form rows, extract them as
  attribute/value pairs.
- For unstructured text, optionally use an Ollama LLM to extract
  [{"attribute": "...", "value": "..."}] pairs.
- For Excel/CSV structured multi-column inputs, preserve all columns and only
  standardize the detected attribute/value columns.
"""

import re
import io
import os
import json
import urllib.request
import urllib.error
import pdfplumber
import pandas as pd
from typing import Union
from docx import Document
import xml.etree.ElementTree as ET


_SKIP_PATTERNS = re.compile(
    r"^(field\s*label|attribute|key|column|header|name|label|field|value|#|no\.?)$",
    re.IGNORECASE,
)

_JSON_ARRAY_PATTERN = re.compile(r"\[\s*\{.*\}\s*\]", re.DOTALL)

ATTRIBUTE_CANDIDATES = {
    "attribute", "raw_attribute", "field", "field_name", "field label",
    "label", "key", "name", "attribute name", "raw attribute"
}

VALUE_CANDIDATES = {
    "value", "raw_value", "field_value", "val", "attribute_value",
    "attribute value", "raw value"
}


def _clean_cell(v) -> str:
    if v is None:
        return ""
    s = str(v).strip()
    if s.lower() in ("nan", "none", ""):
        return ""
    return s


def _normalize_col_name(name: str) -> str:
    s = _clean_cell(name).lower()
    s = re.sub(r"[_\-\s]+", " ", s)
    return s.strip()


def _is_header_row(attr: str, val: str) -> bool:
    attr = (attr or "").strip()
    val = (val or "").strip()
    if not attr:
        return True
    if _SKIP_PATTERNS.match(attr):
        return True
    if attr.lower() in {"field label", "attribute", "name"} and val.lower() in {"value", "values"}:
        return True
    return False


def _row_to_kv(row) -> tuple[str, str]:
    """Return (attr, value) using first non-empty cell as attribute and next non-empty cell as value."""
    cells = [_clean_cell(c) for c in row if _clean_cell(c)]
    if len(cells) < 2:
        return "", ""
    attr = cells[0]
    value = cells[1]
    return attr, value


def _safe_json_loads(text: str):
    try:
        return json.loads(text)
    except Exception:
        return None


def _extract_json_array(text: str):
    if not text:
        return None

    parsed = _safe_json_loads(text)
    if isinstance(parsed, list):
        return parsed
    if isinstance(parsed, dict):
        if isinstance(parsed.get("records"), list):
            return parsed["records"]
        if isinstance(parsed.get("items"), list):
            return parsed["items"]
        if isinstance(parsed.get("data"), list):
            return parsed["data"]
        if isinstance(parsed.get("attributes"), list):
            return parsed["attributes"]

    match = _JSON_ARRAY_PATTERN.search(text)
    if not match:
        return None

    candidate = match.group(0)
    parsed = _safe_json_loads(candidate)
    if isinstance(parsed, list):
        return parsed
    return None


def _dedupe_records(records: list[dict]) -> list[dict]:
    cleaned = []
    seen = set()

    for rec in records:
        attr = str(rec.get("attribute", "")).strip()
        val = str(rec.get("value", "")).strip()

        if not attr or not val:
            continue
        if _is_header_row(attr, val):
            continue
        if re.fullmatch(r"[\d\s.,]+", attr):
            continue

        key = (attr.lower(), val)
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(rec)

    return cleaned


def _looks_like_header_row(row_values: list[str]) -> bool:
    vals = [_clean_cell(v) for v in row_values]
    non_empty = [v for v in vals if v]
    if not non_empty:
        return False

    normalized = {_normalize_col_name(v) for v in non_empty}
    if normalized & ATTRIBUTE_CANDIDATES and normalized & VALUE_CANDIDATES:
        return True

    alpha_like = sum(1 for v in non_empty if re.fullmatch(r"[A-Za-z0-9 _\-/()%.#]+", v))
    if len(non_empty) >= 3 and alpha_like == len(non_empty):
        shortish = sum(1 for v in non_empty if len(v) <= 30)
        if shortish == len(non_empty):
            return True

    return False


def _make_unique_headers(headers: list[str]) -> list[str]:
    out = []
    counts = {}
    for i, h in enumerate(headers):
        base = _clean_cell(h) or f"col_{i+1}"
        if base not in counts:
            counts[base] = 1
            out.append(base)
        else:
            counts[base] += 1
            out.append(f"{base}_{counts[base]}")
    return out


def _prepare_structured_df(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy().fillna("")
    df = df.apply(lambda col: col.map(_clean_cell))
    df = df[df.apply(lambda row: any(_clean_cell(v) for v in row), axis=1)].reset_index(drop=True)
    if df.empty:
        return df

    first_row = df.iloc[0].tolist()
    if _looks_like_header_row(first_row):
        headers = _make_unique_headers(first_row)
        data = df.iloc[1:].reset_index(drop=True).copy()
        data.columns = headers
        return data

    headers = [f"col_{i+1}" for i in range(len(df.columns))]
    df.columns = headers
    return df.reset_index(drop=True)


def _detect_attr_value_columns(columns: list[str]) -> tuple[str, str]:
    normalized_map = {col: _normalize_col_name(col) for col in columns}

    attr_col = None
    value_col = None

    for col, norm in normalized_map.items():
        if norm in ATTRIBUTE_CANDIDATES:
            attr_col = col
            break

    for col, norm in normalized_map.items():
        if norm in VALUE_CANDIDATES:
            value_col = col
            break

    if attr_col is None and columns:
        attr_col = columns[0]
    if value_col is None and len(columns) >= 2:
        value_col = columns[1]
    elif value_col is None and columns:
        value_col = columns[0]

    return attr_col, value_col


def dataframe_to_records_preserve_columns(df: pd.DataFrame) -> list[dict]:
    """
    Convert a structured sheet/dataframe into records while preserving all columns.
    Adds standardized keys:
      - attribute
      - value
    and leaves all original columns intact.
    """
    prepared = _prepare_structured_df(df)
    if prepared.empty:
        return []

    attr_col, value_col = _detect_attr_value_columns(list(prepared.columns))
    records = []
    seen = set()

    for _, row in prepared.iterrows():
        row_dict = {str(col): _clean_cell(val) for col, val in row.to_dict().items()}
        attr = _clean_cell(row_dict.get(attr_col, ""))
        val = _clean_cell(row_dict.get(value_col, ""))

        if not attr or not val:
            continue
        if _is_header_row(attr, val):
            continue
        if re.fullmatch(r"[\d\s.,]+", attr):
            continue

        dedupe_key = (attr.lower(), val, tuple(sorted(row_dict.items())))
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)

        rec = dict(row_dict)
        rec["attribute"] = attr
        rec["value"] = val
        records.append(rec)

    return records


def _simple_free_text_fallback(text: str) -> list[dict]:
    """
    Fallback if LLM is unavailable.
    Uses safe regex rules for common fields.
    """
    records = []
    

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        m = re.match(r"^([A-Za-z][^:=\-]{1,80}?)[\s]*[:=\-]+[\s]*(.*)$", line)
        if m:
            attr = m.group(1).strip()
            val = m.group(2).strip()
            records.append({"attribute": attr, "value": val})

    patterns = [
        (r"\btotal billed amount\b.*?(\d+(?:\.\d+)?\s*rupees?)", "Total Order Value"),
        (r"\btotal amount\b.*?(\d+(?:\.\d+)?\s*rupees?)", "Total Order Value"),
        (r"\bamount\b.*?(\d+(?:\.\d+)?\s*rupees?)", "Total Order Value"),
        (r"\bfor\s+(\d+)\s+(?:units?|pcs?|pieces?|items?)\b", "Quantity"),
        (r"\bon\s+(\d{1,2}\s+[A-Za-z]+\s+\d{4})\b", "Order Date"),
        (r"\bpayment.*?\b(UPI|Cash|Card|Net Banking)\b", "Payment Method"),
        (r"\bdelivery city is\s+([A-Za-z ]+?)(?:\s+and|\.$)", "Address"),
        (r"\border reference\s+is\s+([A-Z0-9\-]+)", "Order Number"),
    ]

    for pattern, attr in patterns:
        for m in re.finditer(pattern, text, flags=re.IGNORECASE):
            value = m.group(1).strip()
            records.append({"attribute": attr, "value": value})

    m = re.search(r"^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\s+placed an order", text)
    if m:
        records.append({"attribute": "Customer Name", "value": m.group(1).strip()})

    m = re.search(
        r"\bfor\s+\d+\s+(?:units?|pcs?|pieces?|items?)\s+of\s+([A-Za-z0-9 ]+?)(?:\.| at | worth )",
        text,
        flags=re.IGNORECASE,
    )
    if m:
        records.append({"attribute": "Product Name", "value": m.group(1).strip()})

    return _dedupe_records(records)


def is_likely_unstructured_text(text: str) -> bool:
    """
    Heuristic check:
    - long paragraph-like lines
    - low key:value density
    - natural language sentence structure
    """
    if not text or not text.strip():
        return False

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return False

    text_stripped = text.strip()
    kv_like_lines = sum(
        1 for ln in lines
        if re.search(r"[:=\-]", ln) and len(ln.split()) <= 12
    )
    avg_words = sum(len(ln.split()) for ln in lines) / max(len(lines), 1)
    long_lines = sum(1 for ln in lines if len(ln.split()) >= 10)
    sentence_punct = len(re.findall(r"[.?!]", text_stripped))

    if len(lines) <= 5 and avg_words >= 12:
        return True
    if long_lines >= max(2, len(lines) // 2) and kv_like_lines == 0:
        return True
    if sentence_punct >= 2 and kv_like_lines == 0 and avg_words >= 8:
        return True

    return False


def should_use_llm_for_text(text: str, structured_record_count: int) -> bool:
    """
    Use LLM if:
    - there are no/very few structured records
    - and the raw text looks narrative/unstructured
    """
    if not text or not text.strip():
        return False

    if structured_record_count == 0 and is_likely_unstructured_text(text):
        return True

    if structured_record_count <= 1 and is_likely_unstructured_text(text):
        return True

    return False


def extract_raw_text_from_pdf(file: Union[str, io.BytesIO]) -> str:
    chunks = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            if txt.strip():
                chunks.append(txt)
    return "\n".join(chunks).strip()

def extract_raw_text_from_docx(file: Union[str, io.BytesIO]) -> str:
    doc = Document(file)
    parts = []

    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            cells = [_clean_cell(cell.text) for cell in row.cells if _clean_cell(cell.text)]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_raw_text_from_excel(file: Union[str, io.BytesIO]) -> str:
    xl = pd.read_excel(file, sheet_name=None, header=None, dtype=str)
    parts = []

    for _, df in xl.items():
        df = df.fillna("")
        for _, row in df.iterrows():
            cells = [_clean_cell(v) for v in row.tolist() if _clean_cell(v)]
            if not cells:
                continue
            parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_raw_text_from_csv(file: Union[str, io.BytesIO]) -> str:
    df = pd.read_csv(file, header=None, dtype=str).fillna("")
    parts = []

    for _, row in df.iterrows():
        cells = [_clean_cell(v) for v in row.tolist() if _clean_cell(v)]
        if not cells:
            continue
        parts.append(" | ".join(cells))

    return "\n".join(parts).strip()
def _split_text_into_chunks(text: str, max_chars: int = 2200, overlap: int = 250) -> list[str]:
    """
    Split long text into overlapping chunks to avoid timeout and improve recall.
    Tries to split on paragraph/newline/sentence boundaries when possible.
    """
    text = (text or "").strip()
    if not text:
        return []

    if len(text) <= max_chars:
        return [text]

    chunks = []
    start = 0
    n = len(text)

    while start < n:
        end = min(start + max_chars, n)

        if end < n:
            window = text[start:end]

            split_pos = max(
                window.rfind("\n\n"),
                window.rfind("\n"),
                window.rfind(". "),
                window.rfind("; "),
            )

            if split_pos > int(max_chars * 0.6):
                end = start + split_pos + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= n:
            break

        start = max(end - overlap, start + 1)

    return chunks


def _call_ollama_extract(
    chunk_text: str,
    model: str,
    base_url: str,
    timeout: int,
) -> list[dict] | None:
    """
    Single Ollama extraction call for one chunk.
    Returns parsed list[dict] or None if parsing/call fails.
    """
    prompt = f"""
You are an information extraction engine.

Extract ALL explicit attribute-value pairs from the text below.

Return ONLY a JSON array.
Do not return markdown.
Do not return explanations.
Do not return any text before or after the JSON.

Each item must have exactly this structure:
{{
  "attribute": "...",
  "value": "..."
}}

Rules:
1. Extract every clearly stated field from the text.
2. Use natural field names based on the text itself.
3. Do not restrict extraction to a fixed schema.
4. Do not merge multiple different facts into one field.
5. If multiple products, quantities, IDs, dates, or statuses are mentioned separately, extract them separately.
6. Do not invent missing values.
7. Do not skip IDs, dates, names, amounts, statuses, contacts, addresses, or locations.
8. If nothing is found, return [].

Good example:
[
  {{"attribute": "customer name", "value": "John Smith"}},
  {{"attribute": "order reference", "value": "ORD1024"}},
  {{"attribute": "tracking id", "value": "TRK88901"}},
  {{"attribute": "billing entity", "value": "JS Retail Private Limited"}},
  {{"attribute": "currency", "value": "INR"}}
]

TEXT:
{chunk_text}
""".strip()

    payload = json.dumps(
        {
            "model": model,
            "prompt": prompt,
            "stream": False
        }
    ).encode("utf-8")

    req = urllib.request.Request(
        url=f"{base_url}/api/generate",
        data=payload,
        headers={"Content-Type": "application/json"},
        method="POST",
    )

    with urllib.request.urlopen(req, timeout=timeout) as resp:
        response_data = json.loads(resp.read().decode("utf-8", errors="ignore"))
        raw_response = response_data.get("response", "").strip()

    print("Raw LLM response:", raw_response[:1500])

    extracted = _extract_json_array(raw_response)
    print("Parsed extracted:", extracted)

    if extracted is None:
        return None

    records = []
    for item in extracted:
        if not isinstance(item, dict):
            continue

        attr = str(item.get("attribute", "")).strip()
        val = str(item.get("value", "")).strip()

        if attr and val:
            records.append({"attribute": attr, "value": val})

    return records


def _merge_chunk_records(all_records: list[dict]) -> list[dict]:
    """
    Final merge + dedupe across all chunk outputs.
    """
    return _dedupe_records(all_records)


def parse_unstructured_text(
    text: str,
    model: str | None = None,
    base_url: str | None = None,
    # timeout: int = 180,
    timeout: int = 300,
    max_chunk_chars: int = 2200,
    overlap: int = 250,
    max_retries_per_chunk: int = 3,
) -> list[dict]:
    """
    Uses Ollama to extract attribute-value pairs from free text.

    Reliable version:
    - splits large text into chunks
    - retries failed chunks
    - falls back per chunk if needed
    - merges and deduplicates all chunk records
    """
    text = (text or "").strip()
    if not text:
        return []

    base_url = (base_url or os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")).rstrip("/")
    model = model or os.getenv("OLLAMA_EXTRACTION_MODEL", "qwen2.5:3b")

    chunks = _split_text_into_chunks(text, max_chars=max_chunk_chars, overlap=overlap)
    print(f"LLM extraction called | model={model} | chunks={len(chunks)}")

    all_records = []

    for idx, chunk in enumerate(chunks, start=1):
        print(f"Processing chunk {idx}/{len(chunks)} | chars={len(chunk)}")
        print("Chunk preview:", chunk[:400])

        chunk_records = None
        last_error = None

        for attempt in range(1, max_retries_per_chunk + 1):
            try:
                print(f"  Attempt {attempt} for chunk {idx}")
                chunk_records = _call_ollama_extract(
                    chunk_text=chunk,
                    model=model,
                    base_url=base_url,
                    timeout=timeout,
                )
                if chunk_records is not None:
                    print(f"  Chunk {idx} success | extracted={len(chunk_records)}")
                    break
            except Exception as e:
                last_error = e
                print(f"  Chunk {idx} failed on attempt {attempt}: {str(e)}")

        if chunk_records is None:
            print(f"Chunk {idx} using fallback due to failure: {str(last_error) if last_error else 'parse failure'}")
            chunk_records = _simple_free_text_fallback(chunk)

        all_records.extend(chunk_records)

    merged = _merge_chunk_records(all_records)
    print("Final merged records count:", len(merged))
    print("Final merged records preview:", merged[:20])

    return merged

# def parse_pdf(file: Union[str, io.BytesIO]) -> tuple[list[dict], str]:
#     records = []
#     doc_type = "keyvalue"

#     with pdfplumber.open(file) as pdf:
#         all_text = []
#         seen = set()

#         for page in pdf.pages:
#             tables = page.extract_tables() or []
#             page_had_table = False

#             for table in tables:
#                 if not table:
#                     continue

#                 page_had_table = True
#                 clean_rows = [row for row in table if row and any(_clean_cell(c) for c in row)]

#                 for row in clean_rows:
#                     attr, val = _row_to_kv(row)
#                     if not attr or _is_header_row(attr, val):
#                         continue
#                     if re.fullmatch(r"[\d\s.,]+", attr):
#                         continue

#                     key = (attr.lower(), val)
#                     if key in seen:
#                         continue
#                     seen.add(key)
#                     records.append({"attribute": attr, "value": val})

#             text = page.extract_text() or ""
#             if text.strip():
#                 all_text.append(text)

#         if not records:
#             for line in "\n".join(all_text).splitlines():
#                 line = line.strip()
#                 if not line or len(line) < 3:
#                     continue

#                 m = re.match(r"^([A-Za-z][^:=\-]{1,80}?)[\s]*[:=\-]+[\s]*(.*)$", line)
#                 if m:
#                     attr = m.group(1).strip()
#                     val = m.group(2).strip()
#                     if _is_header_row(attr, val):
#                         continue
#                     records.append({"attribute": attr, "value": val})

#         if not records and is_likely_unstructured_text("\n".join(all_text)):
#             doc_type = "unstructured"

#     return records, doc_type
def _looks_like_bad_pdf_table(records: list[dict]) -> bool:
    if not records:
        return False

    suspicious_attr = 0
    suspicious_val = 0
    sample = records[:20]

    for rec in sample:
        attr = str(rec.get("attribute", "")).strip()
        val = str(rec.get("value", "")).strip()

        if re.fullmatch(r"\d+", attr) or len(attr) <= 2:
            suspicious_attr += 1

        if val.count("=") >= 2:
            suspicious_val += 1

    if suspicious_attr >= max(2, len(sample) // 3):
        return True

    if suspicious_val >= max(2, len(sample) // 3):
        return True

    return False


def parse_pdf(file: Union[str, io.BytesIO]) -> tuple[list[dict], str]:
    records = []
    doc_type = "keyvalue"

    with pdfplumber.open(file) as pdf:
        all_text = []
        seen = set()

        for page in pdf.pages:
            tables = page.extract_tables() or []

            for table in tables:
                if not table:
                    continue

                clean_rows = [row for row in table if row and any(_clean_cell(c) for c in row)]

                for row in clean_rows:
                    attr, val = _row_to_kv(row)
                    if not attr or _is_header_row(attr, val):
                        continue
                    if re.fullmatch(r"[\d\s.,]+", attr):
                        continue

                    key = (attr.lower(), val)
                    if key in seen:
                        continue
                    seen.add(key)
                    records.append({"attribute": attr, "value": val})

            text = page.extract_text() or ""
            if text.strip():
                all_text.append(text)

        raw_text = "\n".join(all_text).strip()

        # If table rows look broken, ignore them
        if _looks_like_bad_pdf_table(records):
            records = []

        # Try key:value line parsing only if no good records
        if not records:
            for line in raw_text.splitlines():
                line = line.strip()
                if not line or len(line) < 3:
                    continue

                m = re.match(r"^([A-Za-z][^:=\-]{1,80}?)[\s]*[:=\-]+[\s]*(.*)$", line)
                if m:
                    attr = m.group(1).strip()
                    val = m.group(2).strip()
                    if _is_header_row(attr, val):
                        continue
                    records.append({"attribute": attr, "value": val})

        # Final routing decision
        if is_likely_unstructured_text(raw_text):
            # if not records:
            records = []
            doc_type = "unstructured"
        

    return records, doc_type


def parse_pdf_as_dataframes(file: Union[str, io.BytesIO]) -> list[pd.DataFrame]:
    dfs = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables() or []:
                if not table:
                    continue
                rows = [[_clean_cell(c) for c in row] for row in table if row and any(_clean_cell(c) for c in row)]
                if len(rows) >= 2:
                    header = [c or f"Col{i}" for i, c in enumerate(rows[0])]
                    dfs.append(pd.DataFrame(rows[1:], columns=header))
    return dfs

def parse_docx(file: Union[str, io.BytesIO]) -> tuple[list[dict], str]:
    records = []
    seen = set()
    doc_type = "keyvalue"

    doc = Document(file)
    all_text = []

    # Table extraction first
    for table in doc.tables:
        for row in table.rows:
            cells = [_clean_cell(cell.text) for cell in row.cells if _clean_cell(cell.text)]
            if len(cells) >= 2:
                attr = cells[0]
                val = cells[1]

                if not attr or _is_header_row(attr, val):
                    continue
                if re.fullmatch(r"[\d\s.,]+", attr):
                    continue

                key = (attr.lower(), val)
                if key in seen:
                    continue
                seen.add(key)
                records.append({"attribute": attr, "value": val})

    # Paragraph text
    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if txt:
            all_text.append(txt)

    raw_text = "\n".join(all_text).strip()

    # If no structured rows, try key:value line parsing
    if not records:
        for line in raw_text.splitlines():
            line = line.strip()
            if not line or len(line) < 3:
                continue

            m = re.match(r"^([A-Za-z][^:=\\-]{1,80}?)[\\s]*[:=\\-]+[\\s]*(.*)$", line)
            if m:
                attr = m.group(1).strip()
                val = m.group(2).strip()
                if _is_header_row(attr, val):
                    continue
                records.append({"attribute": attr, "value": val})

    if is_likely_unstructured_text(raw_text):
        doc_type = "unstructured"

    return records, doc_type
def _flatten_json_to_records(data, parent_key="") -> list[dict]:
    records = []

    if isinstance(data, dict):
        for key, value in data.items():
            full_key = f"{parent_key}.{key}" if parent_key else str(key)

            if isinstance(value, (dict, list)):
                records.extend(_flatten_json_to_records(value, full_key))
            else:
                records.append({
                    "attribute":  _friendly_attr_from_path(full_key),
                    "value": "" if value is None else str(value),
                    "raw_attribute": full_key
                })

    elif isinstance(data, list):
        for idx, item in enumerate(data, start=1):
            list_key = f"{parent_key}_{idx}" if parent_key else f"item_{idx}"
            records.extend(_flatten_json_to_records(item, list_key))

    return _dedupe_records(records)


def parse_json(file: Union[str, io.BytesIO]) -> tuple[list[dict], str]:
    file.seek(0)
    raw = file.read() if hasattr(file, "read") else open(file, "rb").read()
    text = raw.decode("utf-8", errors="ignore")

    data = json.loads(text)
    records = _flatten_json_to_records(data)

    return records, "json"


def extract_raw_text_from_json(file: Union[str, io.BytesIO]) -> str:
    raw = file.read() if hasattr(file, "read") else open(file, "rb").read()
    return raw.decode("utf-8", errors="ignore")


def _flatten_xml_element(elem, parent_key="") -> list[dict]:
    records = []
    tag_path = f"{parent_key}.{elem.tag}" if parent_key else elem.tag

    for attr_key, attr_val in elem.attrib.items():
        raw_path = f"{tag_path}.{attr_key}"

        records.append({
            "attribute": _friendly_attr_from_path(raw_path),  # ✅ FIX
            "value": attr_val,
            "raw_attribute": raw_path
        })

    text = (elem.text or "").strip()
    if text:
        records.append({
            "attribute": _friendly_attr_from_path(tag_path),  # ✅ FIX
            "value": text,
            "raw_attribute": tag_path
        })

    for child in list(elem):
        records.extend(_flatten_xml_element(child, tag_path))

    return records



def parse_xml(file: Union[str, io.BytesIO]) -> tuple[list[dict], str]:
    file.seek(0)
    raw = file.read() if hasattr(file, "read") else open(file, "rb").read()
    root = ET.fromstring(raw)

    records = _flatten_xml_element(root)
    return _dedupe_records(records), "xml"


def extract_raw_text_from_xml(file: Union[str, io.BytesIO]) -> str:
    raw = file.read() if hasattr(file, "read") else open(file, "rb").read()
    return raw.decode("utf-8", errors="ignore")

def parse_excel(file: Union[str, io.BytesIO]) -> tuple[dict[str, pd.DataFrame], str]:
    xl = pd.read_excel(file, sheet_name=None, header=None, dtype=str)
    sheets = {}
    for sheet_name, df in xl.items():
        df = df.dropna(how="all").dropna(axis=1, how="all").fillna("")
        df = df.apply(lambda col: col.map(lambda x: str(x).strip()))
        df = df[df.apply(lambda row: any(_clean_cell(v) for v in row), axis=1)].reset_index(drop=True)
        if not df.empty:
            sheets[sheet_name] = df
    return sheets, "keyvalue"



def extract_kv_from_excel_sheet(df: pd.DataFrame) -> list[dict]:
    records = []
    seen = set()

    for _, row in df.iterrows():
        attr, val = _row_to_kv(list(row))
        if not attr or _is_header_row(attr, val):
            continue
        key = (attr.lower(), val)
        if key in seen:
            continue
        seen.add(key)
        records.append({"attribute": attr, "value": val})

    return records


def extract_structured_records_from_excel_sheet(df: pd.DataFrame) -> list[dict]:
    return dataframe_to_records_preserve_columns(df)


def extract_structured_records_from_csv_df(df: pd.DataFrame) -> list[dict]:
    return dataframe_to_records_preserve_columns(df)


def extract_tabular_from_excel_sheet(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    headers = [(_clean_cell(v) or f"Col{i}") for i, v in enumerate(df.iloc[0].tolist())]
    new_df = df.iloc[1:].copy()
    new_df.columns = headers
    new_df = new_df[new_df.apply(lambda row: any(_clean_cell(v) for v in row), axis=1)]
    return new_df.reset_index(drop=True)


def parse_csv(file: Union[str, io.BytesIO]) -> tuple[pd.DataFrame, str]:
    df = pd.read_csv(file, header=None, dtype=str).fillna("")
    df = df[df.apply(lambda row: any(_clean_cell(v) for v in row), axis=1)].reset_index(drop=True)
    return df, "keyvalue"

def _friendly_attr_from_path(path: str) -> str:
    path = str(path or "").strip()
    parts = path.split(".")
    parts = [p for p in parts if p.lower() not in ("data", "root", "record", "records")]

    cleaned_parts = []
    for part in parts:
        part = part.replace("_", " ").replace("-", " ")
        part = re.sub(r"([a-z])([A-Z])", r"\1 \2", part)
        part = part.lower().strip()
        if part:
            cleaned_parts.append(part)

    return " ".join(cleaned_parts).strip()